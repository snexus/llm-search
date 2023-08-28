import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Union

import docx
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from loguru import logger


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class HeadingSequence:
    headings: List[Heading] = field(default_factory=list)

    def add(self, text: str, level: int):
        """Adds new headings to the heading sequence

        Args:
            text (str): text of the heading
            level (int): level of the heading (0 - most significant, e.g. title)
        """

        current_heading = Heading(level, text.strip())

        while self.headings and self.headings[-1].level >= current_heading.level:
            self.headings.pop()
        self.headings.append(current_heading)

    @property
    def path(self) -> str:
        s = [t.text for t in self.headings]
        return "/".join(s)


## Based on langchain's splitter


class RecursiveCharacterTextSplitter:
    """Splitting text by recursively look at characters.

    Recursively tries to split by different characters to find one
    that works.
    """

    def __init__(
        self,
        chunk_size: int,
        length_function: Callable[[str], int] = len,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None,
        keep_separator: bool = True,
        is_separator_regex: bool = False,
        **kwargs: Any,
    ) -> None:
        """Create a new TextSplitter."""
        #        super().__init__(keep_separator=keep_separator, **kwargs)
        self._keep_separator = keep_separator
        self._separators = separators or ["\n\n", "\n", " ", ""]
        self._is_separator_regex = is_separator_regex
        self._chunk_size = chunk_size
        self._length_function = length_function
        self._chunk_overlap = chunk_overlap

    def _split_text(self, text: str, separators: List[str]) -> List[str]:
        """Split incoming text and return chunks."""
        final_chunks = []
        # Get appropriate separator to use
        separator = separators[-1]
        new_separators = []
        for i, _s in enumerate(separators):
            _separator = _s if self._is_separator_regex else re.escape(_s)
            if _s == "":
                separator = _s
                break
            if re.search(_separator, text):
                separator = _s
                new_separators = separators[i + 1 :]
                break

        _separator = separator if self._is_separator_regex else re.escape(separator)
        splits = _split_text_with_regex(text, _separator, self._keep_separator)

        # Now go merging things, recursively splitting longer texts.
        _good_splits = []
        _separator = "" if self._keep_separator else separator
        for s in splits:
            if self._length_function(s) < self._chunk_size:
                _good_splits.append(s)
            else:
                if _good_splits:
                    merged_text = self._merge_splits(_good_splits, _separator)
                    final_chunks.extend(merged_text)
                    _good_splits = []
                if not new_separators:
                    final_chunks.append(s)
                else:
                    other_info = self._split_text(s, new_separators)
                    final_chunks.extend(other_info)
        if _good_splits:
            merged_text = self._merge_splits(_good_splits, _separator)
            final_chunks.extend(merged_text)
        return final_chunks

    def split_text(self, text: str) -> List[str]:
        return self._split_text(text, self._separators)

    def _join_docs(self, docs: List[str], separator: str) -> Optional[str]:
        text = separator.join(docs)
        text = text.strip()
        if text == "":
            return None
        else:
            return text

    def _merge_splits(self, splits: Iterable[str], separator: str) -> List[str]:
        # We now want to combine these smaller pieces into medium size
        # chunks to send to the LLM.
        separator_len = self._length_function(separator)

        docs = []
        current_doc: List[str] = []
        total = 0
        for d in splits:
            _len = self._length_function(d)
            if (
                total + _len + (separator_len if len(current_doc) > 0 else 0)
                > self._chunk_size
            ):
                if total > self._chunk_size:
                    logger.warning(
                        f"Created a chunk of size {total}, "
                        f"which is longer than the specified {self._chunk_size}"
                    )
                if len(current_doc) > 0:
                    doc = self._join_docs(current_doc, separator)
                    if doc is not None:
                        docs.append(doc)
                    # Keep on popping if:
                    # - we have a larger chunk than in the chunk overlap
                    # - or if we still have any chunks and the length is long
                    while total > self._chunk_overlap or (
                        total + _len + (separator_len if len(current_doc) > 0 else 0)
                        > self._chunk_size
                        and total > 0
                    ):
                        total -= self._length_function(current_doc[0]) + (
                            separator_len if len(current_doc) > 1 else 0
                        )
                        current_doc = current_doc[1:]
            current_doc.append(d)
            total += _len + (separator_len if len(current_doc) > 1 else 0)
        doc = self._join_docs(current_doc, separator)
        if doc is not None:
            docs.append(doc)
        return docs


def _split_text_with_regex(
    text: str, separator: str, keep_separator: bool
) -> List[str]:
    # Now that we have the separator, split the text
    if separator:
        if keep_separator:
            # The parentheses in the pattern keep the delimiters in the result.
            _splits = re.split(f"({separator})", text)
            splits = [_splits[i] + _splits[i + 1] for i in range(1, len(_splits), 2)]
            if len(_splits) % 2 == 0:
                splits += _splits[-1:]
            splits = [_splits[0]] + splits
        else:
            splits = re.split(separator, text)
    else:
        splits = list(text)
    return [s for s in splits if s != ""]


##This function extracts the tables and paragraphs from the document object
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """

    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_table(table: Table) -> List[dict]:
    """Parses a single table recursively. Supports nested tables.

    Args:
        table (Table): Table element to split

    Returns:
        list: List of dictionaries containing rows of the tables
    """

    headers = get_table_header(table)

    out_list = list()
    current_dict = dict()

    for row in table.rows[1:]:
        for header, cell in zip(headers, row.cells):
            if cell.tables:
                current_dict[header] = []
                current_dict[header].append(cell.text)
            else:
                current_dict[header] = cell.text

            for i, nested_table in enumerate(cell.tables):
                r = parse_table(nested_table)
                if r:
                    current_dict[header].append(r)

        out_list.append(current_dict)
        current_dict = dict()

    return out_list


def get_table_header(table: Table):
    header = tuple([cell.text for cell in table.rows[0].cells])
    return header


def docx_splitter(
    path: Union[str, Path], max_chunk_size: int, **additional_splitter_setting
):
    doc = docx.Document(path)

    hs = HeadingSequence()

    out_chunks = []
    current_chunk = ""

    # Splitter for text paragraphs
    paragraph_splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_chunk_size, separators=["\n\n", "\n", ". ", " "]
    )

    # Splitter for json
    json_splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_chunk_size,
        separators=[r"\{.*?\}", r'".*?"\s*:\s*".*?"', r"\n\n", r"\n", r"\. ", r"\."],
        is_separator_regex=True,
        chunk_overlap=0,
    )

    # Iterate over all elements in the document

    for el in iter_block_items(doc):
        if isinstance(el, Paragraph):
            p_text = el.text

            # If text of the paragraph is empty, continue
            if not p_text.strip():
                continue

            # Store headings for metadata
            # if "title" in el.style.name.lower():
            #    hs.add(p_text, level=0)
            elif "heading" in el.style.name.lower():
                heading_level = int(el.style.name.lower().split(" ")[-1])
                hs.add(p_text, heading_level)

            current_chunk = add_or_split(
                p_text,
                paragraph_splitter,
                current_chunk,
                max_chunk_size,
                out_chunks,
                hs,
                type="paragraph",
            )

        elif isinstance(el, Table):
            t_json = json.dumps(parse_table(el))
            current_chunk = add_or_split(
                t_json,
                json_splitter,
                current_chunk,
                max_chunk_size,
                out_chunks,
                hs,
                type="tabular data in json format.",
            )

    return out_chunks


def add_or_split(
    text,
    splitter,
    current_chunk: str,
    chunk_size: int,
    out_chunks: List[Dict[str, str]],
    hs: HeadingSequence,
    type: str,
) -> str:
    """Adds or splits text to out chunks, together with additional metadata

    Args:
        text (str): Text candidate to add
        splitter (Instance of the splitter): Instance of the splitter
        current_chunk (str): Current buffer (not flushed yet)
        chunk_size (int): Target chunk size
        out_chunks (List[str]): List of all historical chunks
        hs (HeadingSequence): Holding the metadata for the text

    Returns:
        str: current chunk (not flushed)
    """
    additional_metadata = {"type": type}

    # Case 1- length of the next paragraph > chunk_size, flulsh the current and split the next
    if len(text) >= chunk_size:
        out_chunks.append(add_metadata(hs, current_chunk, additional_metadata))
        current_chunk = ""

        for ch in splitter.split_text(text):
            out_chunks.append(add_metadata(hs, ch, additional_metadata))

    # Case 2 - next paragraph is smaller than chunk size, but can't be added to the current chunk
    elif len(current_chunk) + len(text) >= chunk_size:
        out_chunks.append(add_metadata(hs, current_chunk, additional_metadata))

        current_chunk = text

    # Case 3 - current chunk is small enough
    else:
        current_chunk += text

    return current_chunk


def add_metadata(
    hs: HeadingSequence, text: str, additional_metadata: Optional[dict] = None
) -> Dict[str, str]:
    metadata_s = ""

    additional_metadata = {} if additional_metadata is None else additional_metadata
    additional_metadata.update({"topic": f"{hs.path}"})

    for k, v in additional_metadata.items():
        if v:
            metadata_s += f"{k}: {v}\n"

    metadata = f"Metadata applicable to the next chunk of text delimited by five stars:\n<< METADATA\n{metadata_s}METADATA\n\n"
    t = metadata + "*****\n" + text + "\n*****"
    chunk = {"text": t, "metadata": {"heading": hs.path}}
    return chunk


if __name__ == "__main__":
    fn = "/home/lacpd1/projects/doc-parsing/doc1.docx"
    parsed = docx_splitter(fn, max_chunk_size=1024)

    for p in parsed:
        print("-------------")
        print(len(p["text"]))
        print(p["text"])
